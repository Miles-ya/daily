from __future__ import annotations

import hashlib
import io
import re
import time
from datetime import datetime
from urllib.parse import urljoin, urlsplit, urlunsplit

import requests
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader

from daily.collectors.base import DiscoveredItem
from daily.models import PolicyDocument, utc_now
from daily.pipeline.policy import classify_status, is_policy_candidate


DATE_RE = re.compile(r"(20\d{2})[年./-](\d{1,2})[月./-](\d{1,2})日?")
DATE_IN_URL_RE = re.compile(r"/t(20\d{2})(\d{2})(\d{2})_")
TIME_RE = re.compile(r"(?:发布时间|发布日期|时间)[：:\s]*(?:20\d{2}[年./-]\d{1,2}[月./-]\d{1,2}日?)?\s*(\d{1,2}:\d{2})")
DOC_NUMBER_RE = re.compile(r"([\u4e00-\u9fff]{1,12}(?:字|发|办|令|公告)?〔20\d{2}〕\d+号)")
EFFECTIVE_RE = re.compile(r"自(20\d{2})年(\d{1,2})月(\d{1,2})日起(?:施行|实施)")
CONTENT_SELECTORS = ".article, .article-content, .content, .TRS_Editor, .pages_content, .view, .detail, #UCAP-CONTENT, #zoom, main"


def canonicalize_policy_url(url: str) -> str:
    parts = urlsplit(url)
    path = re.sub(r"/{2,}", "/", parts.path)
    return urlunsplit((parts.scheme.lower(), parts.netloc.lower(), path, "", ""))


class GenericPolicyCollector:
    def __init__(self, source_id: str, config: dict, timeout: int = 25, retries: int = 3,
                 user_agent: str = "PolicyRadar/0.1"):
        self.source_id = source_id
        self.source_name = config["name"]
        self.list_urls = config["list_urls"]
        self.patterns = tuple(config.get("include_patterns", ()))
        self.max_documents = config.get("max_documents", 30)
        self.timeout = timeout
        self.retries = retries
        self.session = requests.Session()
        self.session.headers.update({"User-Agent": user_agent, "Accept-Language": "zh-CN,zh;q=0.9"})

    def _get(self, url: str) -> requests.Response:
        last_error: Exception | None = None
        for attempt in range(self.retries):
            try:
                response = self.session.get(url, timeout=self.timeout)
                response.raise_for_status()
                if "text" in response.headers.get("content-type", "") or response.url.lower().endswith((".htm", ".html", "/")):
                    response.encoding = response.apparent_encoding or "utf-8"
                return response
            except requests.RequestException as exc:
                last_error = exc
                if attempt + 1 < self.retries:
                    time.sleep(0.5 * (2 ** attempt))
        raise RuntimeError(f"请求失败: {self.source_name}: {last_error}")

    def discover(self) -> list[DiscoveredItem]:
        found: dict[str, DiscoveredItem] = {}
        errors: list[str] = []
        for list_url in self.list_urls:
            try:
                soup = BeautifulSoup(self._get(list_url).text, "html.parser")
            except Exception as exc:
                errors.append(str(exc))
                continue
            for anchor in soup.select("a[href]"):
                title = anchor.get_text(" ", strip=True)
                if not title or len(title) < 6:
                    continue
                url = canonicalize_policy_url(urljoin(list_url, anchor.get("href", "")))
                if self.patterns and not any(pattern in url for pattern in self.patterns):
                    continue
                if not is_policy_candidate(title):
                    continue
                container = anchor.find_parent(["li", "tr", "article", "div"])
                nearby = container.get_text(" ", strip=True) if container else (anchor.parent.get_text(" ", strip=True) if anchor.parent else title)
                matches = list(DATE_RE.finditer(nearby)) or list(DATE_RE.finditer(url)) or list(DATE_IN_URL_RE.finditer(url))
                match = max(matches, key=lambda value: tuple(map(int, value.groups()))) if matches else None
                publish_date = f"{int(match.group(1)):04d}-{int(match.group(2)):02d}-{int(match.group(3)):02d}" if match else None
                found[url] = DiscoveredItem(url=url, title=title, publish_date=publish_date)
        if not found and errors:
            raise RuntimeError("; ".join(errors))
        return sorted(found.values(), key=lambda item: (item.publish_date or "", item.url), reverse=True)[:self.max_documents]

    @staticmethod
    def _binary_text(response: requests.Response, url: str) -> tuple[str, bool]:
        lowered = url.lower()
        try:
            if lowered.endswith(".pdf") or "pdf" in response.headers.get("content-type", ""):
                return "\n".join(page.extract_text() or "" for page in PdfReader(io.BytesIO(response.content)).pages), True
            if lowered.endswith(".docx"):
                doc = DocxDocument(io.BytesIO(response.content))
                return "\n".join(paragraph.text for paragraph in doc.paragraphs), True
        except Exception:
            return "", False
        return "", not lowered.endswith(".doc")

    def collect(self, item: DiscoveredItem) -> PolicyDocument:
        response = self._get(item.url)
        direct_text, complete = self._binary_text(response, item.url)
        attachments: list[dict[str, str]] = []
        if direct_text:
            title = item.title or item.url.rsplit("/", 1)[-1]
            page_text = direct_text
            content = direct_text
        else:
            soup = BeautifulSoup(response.text, "html.parser")
            title_node = soup.select_one("h1, .title, .article-title")
            page_title = soup.title.get_text(" ", strip=True) if soup.title else ""
            title = (title_node.get_text(" ", strip=True) if title_node else "") or item.title or page_title
            content_node = soup.select_one(CONTENT_SELECTORS) or soup.body
            if content_node is None:
                raise ValueError(f"无法定位正文: {item.url}")
            for node in content_node.select("script, style, nav, .share"):
                node.decompose()
            content = "\n".join(line.strip() for line in content_node.get_text("\n").splitlines() if line.strip())
            page_text = soup.get_text(" ", strip=True)
            for anchor in content_node.select("a[href]"):
                target = canonicalize_policy_url(urljoin(item.url, anchor.get("href", "")))
                if not target.lower().endswith((".pdf", ".doc", ".docx")):
                    continue
                attachment = {"title": anchor.get_text(" ", strip=True) or target.rsplit("/", 1)[-1], "url": target}
                attachments.append(attachment)
                if len(content) < 50000:
                    try:
                        extra, extracted = self._binary_text(self._get(target), target)
                        complete = complete and extracted
                        if extra:
                            content += f"\n\n附件：{attachment['title']}\n{extra}"
                    except Exception:
                        complete = False
        if not title or len(content) < 20:
            raise ValueError(f"正文过短或缺少标题: {item.url}")
        publish_date = item.publish_date
        if not publish_date:
            candidates = list(DATE_RE.finditer(page_text[:5000]))
            if candidates:
                today = datetime.now().date()
                valid = [value for value in candidates if datetime(int(value.group(1)), int(value.group(2)), int(value.group(3))).date() <= today]
                date_match = max(valid or candidates, key=lambda value: tuple(map(int, value.groups())))
                publish_date = f"{int(date_match.group(1)):04d}-{int(date_match.group(2)):02d}-{int(date_match.group(3)):02d}"
        time_match = TIME_RE.search(page_text)
        number_match = DOC_NUMBER_RE.search(page_text[:4000])
        effective_match = EFFECTIVE_RE.search(content)
        effective_date = None
        if effective_match:
            effective_date = f"{int(effective_match.group(1)):04d}-{int(effective_match.group(2)):02d}-{int(effective_match.group(3)):02d}"
        canonical = canonicalize_policy_url(item.url)
        content = content[:60000]
        return PolicyDocument(
            id=hashlib.sha256(canonical.encode()).hexdigest()[:20], source_id=self.source_id,
            source_name=self.source_name, url=item.url, canonical_url=canonical, title=title,
            publish_date=publish_date, publish_time=time_match.group(1) if time_match else None,
            crawl_time=utc_now(), content=content, content_hash=hashlib.sha256(content.encode()).hexdigest(),
            issuing_bodies=[self.source_name], document_number=number_match.group(1) if number_match else "",
            effective_date=effective_date, policy_status=classify_status(title), attachments=attachments,
            extraction_complete=complete, raw_metadata={"http_status": response.status_code, "content_type": response.headers.get("content-type", "")},
        )
